from pathlib import Path
from datetime import datetime
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE


OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)


MASTER_PROFILE = {
    "name": "Paulo Henrique Santos Oliveira",
    "phone": "(71) 99349-4443",
    "email": "henriqueoliveirarh93@gmail.com",
    "linkedin": "https://www.linkedin.com/in/paulo-oliveira-933a9254/",
    "location": "Salvador/BA",

    "headline": (
        "Recursos Humanos | Departamento Pessoal | Gente & Gestão | "
        "Power BI | Indicadores"
    ),

    "summary": (
        "Especialista em Recursos Humanos Generalista & Estratégico, com forte "
        "foco em Data-Driven RH e Tech Recruiting. Mais de 10 anos de experiência "
        "na Coordenação de RH, Administração de Pessoal, R&S, T&D e Gestão "
        "Trabalhista. Experiência comprovada em implantação de setores de RH, "
        "otimização de custos, redução de turnover e horas extras e conformidade "
        "legal, incluindo ISO 9001 e e-Social. Proficiência em Power BI, "
        "dashboards e sistemas como RM TOTVS, Protheus e SAP."
    ),

    "skills": [
        "Recursos Humanos",
        "Departamento Pessoal",
        "Administração de Pessoal",
        "Folha de Pagamento",
        "Recrutamento e Seleção",
        "Treinamento e Desenvolvimento",
        "Gestão de Equipes",
        "Avaliação de Desempenho",
        "Indicadores e KPIs",
        "Power BI",
        "Dashboards",
        "Excel Avançado",
        "Visual Basic / VBA",
        "RM TOTVS / Protheus",
        "RM Labore",
        "SAP",
        "e-Social",
        "Legislação Trabalhista",
        "Relações Trabalhistas",
        "ISO 9001",
    ],

    "experiences": [
        {
            "company": "Logic Soluções Logísticas",
            "role": "Coordenador de Departamento Pessoal",
            "period": "Setembro de 2025 – Atual",
            "bullets": [
                "Liderança do projeto de migração da folha de pagamento da contabilidade externa para operação interna, incluindo importação sistêmica, parametrização de eventos, revisão de encargos e integração contábil.",
                "Redução estimada de 18% nos custos administrativos e aumento do controle estratégico das informações trabalhistas.",
                "Implantação de modelo de governança com padronização de fluxos, matriz de responsabilidades e SLAs operacionais.",
                "Implantação de controle avançado de ponto eletrônico com reestruturação de regras sistêmicas.",
                "Redução de 35% nas inconsistências de batidas e erros em folha.",
                "Implementação de sistema de premiação vinculado à performance e assiduidade, reduzindo o absenteísmo em 12%.",
                "Monitoramento estratégico de horas extras, com redução de 20% no custo mensal.",
                "Estruturação de provisões de férias, 13º salário e encargos, reduzindo variações orçamentárias em 15%.",
                "Gestão de rateios por centro de custo, assegurando precisão contábil e apoio à tomada de decisão financeira.",
                "Desenvolvimento de dashboards gerenciais em Power BI para a Diretoria.",
            ],
        },
        {
            "company": "TPC Logística",
            "role": "Supervisor de Recursos Humanos",
            "period": "Julho de 2022 – Agosto de 2024",
            "bullets": [
                "Liderança da implantação do Centro de Distribuição da Avon, incluindo recrutamento e seleção de 400 colaboradores em apenas dois meses.",
                "Gestão de processos de RH, Gente & Gestão, recrutamento, treinamento e desenvolvimento.",
                "Implantação de controle de horas extras abaixo de 30 minutos via Power BI, resultando em redução de 25% nessas horas extras e economia estimada superior a R$ 8.000,00 por mês.",
                "Criação e implementação de programa social de recrutamento exclusivo para mulheres vítimas de violência doméstica.",
                "Coordenação do time de Gente & Gestão, com implementação de programa contínuo de T&D da liderança.",
                "Implementação de ciclos de avaliação de desempenho.",
                "Apresentação de indicadores e KPIs para a Diretoria utilizando Power BI, RM Labore e TOTVS.",
            ],
        },
        {
            "company": "Inovar Telecom",
            "role": "Supervisor Administrativo e Contratual (PJ)",
            "period": "Outubro de 2024 – Setembro de 2025",
            "bullets": [
                "Liderança na criação e implementação de dashboards gerenciais e KPIs em Power BI para toda a operação.",
                "Criação de visualização de dados de RH que anteriormente não existia na operação.",
                "Desenvolvimento de controles em Excel e Visual Basic (macros), reduzindo em 30% o tempo gasto em rotinas de acompanhamento documental.",
                "Administração e supervisão de rotinas contratuais, acompanhamento documental e emissão de relatórios gerenciais.",
                "Monitoramento e otimização de rotas, obtendo redução de 12% nos custos logísticos da frota.",
            ],
        },
        {
            "company": "Luandre Temporários",
            "role": "Analista de Administração de Pessoal Pleno",
            "period": "Dezembro de 2021 – Junho de 2022",
            "bullets": [
                "Responsável pela Administração de Pessoal de clientes de grande porte, incluindo Mercado Livre, Amazon e Sephora.",
                "Gestão de alto volume de admissões e demissões.",
                "Promoção a Analista Pleno após atuação como Focal Point, com otimização de 35% no tempo de resposta e solução de dúvidas dos colaboradores.",
                "Execução de processos de Administração de Pessoal, incluindo admissão, demissão, contratos e benefícios.",
                "Utilização do sistema RM TOTVS / Protheus.",
            ],
        },
        {
            "company": "ASM – Associação Saúde em Movimento",
            "role": "Analista de Recursos Humanos",
            "period": "Agosto de 2020 – Outubro de 2021",
            "bullets": [
                "Liderança na implantação do setor de RH em novos contratos hospitalares decorrentes de licitações em unidades no RJ, SP, DF, TO e BA.",
                "Responsável pela primeira integração da folha de pagamento com o e-Social, garantindo adequação legal.",
                "Atuação em relações trabalhistas e parceria com o departamento jurídico.",
                "Atuação como preposto em demandas trabalhistas e realização de cálculos de acordos trabalhistas.",
            ],
        },
        {
            "company": "Souza e Filhos",
            "role": "Coordenador de Recursos Humanos",
            "period": "Janeiro de 2019 – Agosto de 2020",
            "bullets": [
                "Responsável pela implantação e adequação dos parâmetros da ISO 9001, contribuindo para a certificação de qualidade.",
                "Redução de 27% no turnover da empresa.",
                "Obtenção de redução de 15% nos custos operacionais por meio de indicadores de RH.",
                "Implantação e internalização da folha de pagamento no sistema Domínio.",
                "Criação e implementação de programas de Trainee, Estágio, Jovem Aprendiz e PNE.",
            ],
        },
    ],

    "education": [
        (
            "Pós-graduação em Power BI: Big Data, Business Intelligence e Analytics",
            "Estácio",
            "Em conclusão – 2026",
        ),
        (
            "Pós-graduação em Gestão Estratégica",
            "Uniasselvi",
            "2021",
        ),
        (
            "Ciência da Computação",
            "UniRuy",
            "4º semestre – Em curso",
        ),
        (
            "Administração de Empresas",
            "Estácio",
            "2020",
        ),
        (
            "Recursos Humanos",
            "Estácio",
            "2015",
        ),
    ],

    "languages": [
        "Espanhol – Leitura avançada | Escrita média | Fala média",
        "Inglês – Leitura média | Escrita média | Fala iniciante",
        "Mandarim – Básico | Em andamento",
    ],
}


def _add_heading(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(5)

    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(31, 78, 121)

    return p


def _add_bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)

    run = p.add_run(text)
    run.font.size = Pt(10)

    return p


def _add_body(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)

    run = p.add_run(text)
    run.font.size = Pt(10)

    return p


def _unique(items: list[str]) -> list[str]:
    return list(dict.fromkeys(item.strip() for item in items if item.strip()))


def _ordered_experiences(resume: dict[str, Any]) -> list[dict[str, Any]]:
    """Use only the authenticated candidate's experiences when supplied."""
    supplied = resume.get("experiences")
    if not supplied:
        return MASTER_PROFILE["experiences"]

    normalized = []
    for item in supplied:
        bullets = item.get("bullets") or [
            line.strip()
            for line in item.get("description", "").splitlines()
            if line.strip()
        ]
        period = item.get("period") or " - ".join(
            value
            for value in (item.get("start_date", ""), item.get("end_date", ""))
            if value
        )
        normalized.append(
            {
                "company": item.get("company", ""),
                "role": item.get("role", ""),
                "period": period,
                "bullets": bullets,
            }
        )
    return normalized


def _configure_document(doc: Document, candidate_name: str):
    section = doc.sections[0]

    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)

    # Cabeçalho / rodapé
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = footer.add_run(f"{candidate_name} • Currículo Profissional")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)


def generate_docx(payload: dict[str, Any]) -> str:
    """
    Gera currículo DOCX completo.

    O payload pode vir diretamente do endpoint de geração.
    Os dados mestres são preservados e a personalização da vaga
    pode sobrescrever título, headline, resumo e competências.
    """

    resume = payload.get("resume", payload)

    candidate_payload = resume.get("candidate")
    if candidate_payload:
        candidate = {
            "name": candidate_payload.get("name", "Candidato"),
            "phone": candidate_payload.get("phone", ""),
            "email": candidate_payload.get("email", ""),
            "linkedin": candidate_payload.get("linkedin", ""),
            "location": candidate_payload.get("location", ""),
        }
    else:
        candidate = dict(MASTER_PROFILE)

    # Personalização da vaga
    job_title = resume.get(
        "target",
        payload.get("job_title", "Profissional de Recursos Humanos")
    )

    headline = resume.get(
        "headline",
        MASTER_PROFILE["headline"]
    )

    summary = resume.get(
        "summary",
        MASTER_PROFILE["summary"]
    )

    skills = _unique(resume.get(
        "skills",
        MASTER_PROFILE["skills"]
    ))

    doc = Document()
    _configure_document(doc, candidate["name"])

    # ============================================================
    # CABEÇALHO
    # ============================================================

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(candidate["name"])
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)

    run = p.add_run(job_title)
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)

    run = p.add_run(headline)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(70, 70, 70)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)

    contact = (
        f'{candidate["location"]}  |  '
        f'{candidate["phone"]}  |  '
        f'{candidate["email"]}'
    )

    run = p.add_run(contact)
    run.font.size = Pt(9)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)

    run = p.add_run(candidate["linkedin"])
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(31, 78, 121)

    # ============================================================
    # RESUMO
    # ============================================================

    _add_heading(doc, "Resumo Profissional")
    _add_body(doc, summary)

    # ============================================================
    # COMPETÊNCIAS
    # ============================================================

    _add_heading(doc, "Competências Técnicas")

    if skills:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)

        r = p.add_run("Prioridades para a vaga: ")
        r.bold = True
        r.font.size = Pt(9.5)

        r = p.add_run(" • ".join(skills))
        r.font.size = Pt(9.5)

    categories = [] if candidate_payload else [
        (
            "RH & Gestão",
            [
                "Recursos Humanos",
                "Gestão de Equipes",
                "Recrutamento e Seleção",
                "Treinamento e Desenvolvimento",
                "Avaliação de Desempenho",
            ],
        ),
        (
            "Departamento Pessoal",
            [
                "Departamento Pessoal",
                "Administração de Pessoal",
                "Folha de Pagamento",
                "Ponto",
                "e-Social",
                "Legislação Trabalhista",
                "Relações Trabalhistas",
            ],
        ),
        (
            "Dados & Tecnologia",
            [
                "Power BI",
                "Dashboards",
                "Indicadores e KPIs",
                "Excel Avançado",
                "Visual Basic / VBA",
                "RM TOTVS / Protheus",
                "RM Labore",
                "SAP",
            ],
        ),
        (
            "Compliance & Processos",
            [
                "ISO 9001",
                "Gestão de Processos",
                "Redução de Custos",
                "Gestão Trabalhista",
            ],
        ),
    ]

    for category, items in categories:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

        r = p.add_run(f"{category}: ")
        r.bold = True
        r.font.size = Pt(9.5)

        r = p.add_run(" • ".join(items))
        r.font.size = Pt(9.5)

    # ============================================================
    # EXPERIÊNCIA
    # ============================================================

    _add_heading(doc, "Experiência Profissional")

    for exp in _ordered_experiences(resume):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(1)

        r = p.add_run(exp["company"])
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(31, 78, 121)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

        r = p.add_run(exp["role"])
        r.bold = True
        r.font.size = Pt(10)

        r = p.add_run(f"  |  {exp['period']}")
        r.italic = True
        r.font.size = Pt(9)

        for bullet in exp["bullets"]:
            _add_bullet(doc, bullet)

    # ============================================================
    # FORMAÇÃO
    # ============================================================

    _add_heading(doc, "Formação Acadêmica")

    education_items = (
        resume.get("education", [])
        if "education" in resume
        else MASTER_PROFILE["education"]
    )
    for item in education_items:
        if isinstance(item, dict):
            course = item.get("course", "")
            institution = item.get("institution", "")
            period = item.get("period", "")
        else:
            course, institution, period = item
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)

        r = p.add_run(course)
        r.bold = True
        r.font.size = Pt(10)

        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_after = Pt(4)

        r = p.add_run(f"{institution} — {period}")
        r.font.size = Pt(9)

    # ============================================================
    # IDIOMAS
    # ============================================================

    _add_heading(doc, "Idiomas")

    languages = (
        resume.get("languages", [])
        if "languages" in resume
        else MASTER_PROFILE["languages"]
    )
    for language in languages:
        _add_bullet(doc, language)

    # ============================================================
    # SALVAR
    # ============================================================

    safe_title = "".join(
        c if c.isalnum() or c in " -_" else "_"
        for c in job_title
    ).strip()

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    safe_name = "_".join(
        part for part in candidate["name"].split() if part
    ) or "Candidato"
    filename = f"Curriculo_{safe_name}_{safe_title}_{timestamp}.docx"

    output_path = OUTPUT_DIR / filename

    doc.save(output_path)

    return str(output_path)
