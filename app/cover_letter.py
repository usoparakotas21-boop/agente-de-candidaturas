from typing import Any

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)


def generate_cover_letter(
    job_title: str,
    company: str,
    profile: dict[str, Any],
    analysis: dict[str, Any],
    personalization: dict[str, Any],
) -> str:
    candidate_name = profile.get(
        "name",
        "Paulo Henrique Santos Oliveira",
    )

    company_name = company or "empresa"

    strengths = analysis.get(
        "strengths",
        [],
        
    )[:4]
    

    active_areas = personalization.get(
        "active_areas",
        [],
    )[:4]
    area_labels = {
        "rh": "Recursos Humanos",
        "dp": "Departamento Pessoal",
        "recrutamento": "Recrutamento e Seleção",
        "td": "Treinamento e Desenvolvimento",
        "dados": "Power BI, indicadores e análise de dados",
        "gestao": "Gestão de Equipes",
        "custos": "Otimização e Redução de Custos",
        "trabalhista": "Relações e Legislação Trabalhista",
    }

    translated_areas = [
        area_labels.get(area, area)
        for area in active_areas
    ]

    strengths_text = ", ".join(strengths)

    if not strengths_text:
        strengths_text = (
            "gestão de pessoas, Departamento Pessoal "
            "e melhoria de processos"
        )

    areas_text = ", ".join(translated_areas)

    if not areas_text:
        areas_text = "Recursos Humanos"

    letter = (
        f"Prezados(as) responsáveis pelo processo seletivo da {company_name},\n\n"
        f"Tenho interesse na oportunidade de {job_title}. "
        f"Minha trajetória profissional possui forte aderência aos desafios "
        f"da posição, especialmente nas áreas de {areas_text}.\n\n"
        f"Entre os principais pontos relacionados à vaga, destaco minha "
        f"experiência em {strengths_text}. Ao longo da minha carreira, "
        f"atuei na implantação e melhoria de processos de Recursos Humanos, "
        f"gestão de equipes, indicadores e redução de custos.\n\n"
        f"Acredito que minha experiência prática, visão analítica e capacidade "
        f"de estruturar operações podem contribuir para os resultados da "
        f"{company_name}.\n\n"
        f"Permaneço à disposição para uma conversa.\n\n"
        f"Atenciosamente,\n"
        f"{candidate_name}"
    )

    return letter

def generate_cover_letter_docx(
    letter: str,
    company: str,
    job_title: str,
) -> str:
    document = Document()

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title.add_run(
        "CARTA DE APRESENTAÇÃO"
    )
    title_run.bold = True
    title_run.font.size = Pt(15)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle_run = subtitle.add_run(
        f"{job_title} | {company}"
    )
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    document.add_paragraph()

    for paragraph_text in letter.split("\n\n"):
        paragraph = document.add_paragraph(
            paragraph_text
        )
        paragraph.paragraph_format.space_after = Pt(10)

    safe_company = "".join(
        character
        if character.isalnum() or character in " -_"
        else "_"
        for character in company
    ).strip()

    timestamp = datetime.now().strftime(
        "%Y%m%d_%H%M%S"
    )

    filename = (
        f"Carta_Apresentacao_"
        f"{safe_company}_{timestamp}.docx"
    )

    output_path = OUTPUT_DIR / filename

    document.save(output_path)

    return str(output_path)
